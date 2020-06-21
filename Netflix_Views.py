import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from datetime import datetime
from collections import Counter
import re
import os
import sys
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import RGBColor, Pt, Length
from docx.enum.style import WD_STYLE

def requirements():
    os.system('pip install pandas matplotlib python-docx')

def data_cleaning():

    def load():

        def d_parser(x): return datetime.strptime(x, '%d/%m/%Y')
        df = pd.read_csv('NetflixViewingHistory.csv', parse_dates=[
                         'Date'], date_parser=d_parser)

        return df

    df = load()
    # Substituting Part and Volume in Season.
    PaVo = re.compile(r'Part|Volume')
    df['Title'] = [PaVo.sub('Season', t) for t in df['Title']]

    # Removing all movies and specials shows from the table
    # checks whether there is a mention for season
    NumS = re.compile('Season\s[1-9]{0,9}')

    show = Counter()
    for t in df['Title']:
        show[t.split(':')[0]] += 1

    movies = []
    for t in df['Title']:
        try:
            if [len(t.split(':')) <= 2 or len(re.findall(NumS, t.split(':')[1:])) == 0] and show[t.split(':')[0]] < 3:
                movies.append(t)
        except TypeError:
            pass

    Count = [len(df[~df['Title'].isin(movies)]), len(movies)]

    df = df[~df['Title'].isin(movies)]
    # After getting all the movies/one time specials (where there is no trace of season number) into 'movies'

    df['Name'] = [t.split(":")[0] for t in df['Title']]  # name column
    df['Description'] = [t.split(":")[-1]
                         for t in df['Title']]  # Description column
    df['DayOfWeek'] = df['Date'].dt.day_name()  # weekday column
    df['Month'] = df['Date'].dt.month_name()  # month column
    df['Year'] = df['Date'].dt.year  # year column

    return df, Count


def Netflix_time(df):
    pattern = re.compile('[0-9]{0,9}\sdays')
    NetflixTime = df['Date'].max() - df['Date'].min()
    Netflix_Days = re.findall(pattern, str(NetflixTime))[0]

    return str(Netflix_Days)

# Finds out the number of shows you ditched after one episode.
# There is must be a better way to do it, but for now this works.
def left(df):

    show = Counter()
    for t in df['Name']:
        show[t] += 1

    ditched = []
    for t in df['Name']:
        if show[t] == 1:
            ditched.append(t)

    return len(ditched)

def views_by_day(df):

    day_grp = df.groupby(['DayOfWeek'])
    weekDays = ['Sunday', 'Monday', 'Tuesday',
                'Wednesday', 'Thursday', 'Friday', 'Saturday']
    day_vals = [day_grp.get_group(day).count()[0] for day in weekDays]

    plt.style.use('seaborn-poster')

    plt.bar(weekDays, day_vals, width=0.45, align='center',
            color='#A61919', edgecolor='#FFFFFF')
    plt.title('Total Views By Day')
    plt.ylabel('Views')

    plt.tight_layout()

    plt.savefig('View_day.jpeg')


def create_pie(Count):
    
    plt.clf()
    plt.style.use("seaborn-poster")

    labels = ['Series', 'Movies']
    # in order to emphasize one piece on the chart.
    explode = [0, 0.1]

    plt.pie(Count, labels=labels, colors=(['#444444', '#B21010']), explode=explode, shadow=True, autopct='%1.1f%%',
            startangle=90, wedgeprops={'edgecolor': 'black'})

    # plt.tight_layout()
    plt.savefig('PieChart.jpeg')

def ten(df):

    plt.clf()
    plt.style.use("seaborn-poster")
        
    table = df.groupby(['Name'], as_index=True).agg(
        Episodes=('Title', 'count'))
    table = table['Episodes'].nlargest(10)
    table = table.reset_index()

    names = [i for i in table['Name']]
    eps = [j for j in table['Episodes']]

    names.reverse()
    eps.reverse()

    plt.barh(names, eps, height=0.7, color='#B21018', edgecolor='black')

    plt.title('10 Longest Shows')
    
    plt.xlabel('Number Of Episodes')
    
    plt.tight_layout()
    
    plt.savefig('10shows.jpeg')
    
    
def create_doc():

    document = Document()
    styles = document.styles

    document.add_heading('Netflix Analysis', 0)

    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(12)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(13)

    paragraph = document.add_paragraph()
    paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    p = document.add_paragraph()
    p.style = 'List Bullet'
    paragraph_format.line_spacing = 1.75
    p.add_run('You have been on Netflix for ')
    p.add_run(Netflix_time(df)).bold = True
    p.add_run('.')

    p = document.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('You have seen ')
    p.add_run(str(Count[1])).bold = True
    p.add_run(' movies.')

    p = document.add_paragraph()
    p.style = 'List Bullet'
    paragraph_format.line_spacing = 1.75
    p.add_run('The most amount of episodes you saw in one day is ')
    p.add_run(str(df.groupby('Date').count().max()['Title'])).bold = True
    p.add_run(' episodes.')

    p = document.add_paragraph()
    p.style = 'List Bullet'
    paragraph_format.line_spacing = 1.75
    p.add_run('You see on average ')
    p.add_run(
        str(round(df.groupby('Date').count().mean()['Title'], 2))).bold = True
    p.add_run(' episodes per day.')

    p = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.style = 'List Bullet'
    paragraph_format.line_spacing = 1.75
    p.add_run('You have ditched ')
    p.add_run(str(left(df))).bold = True
    p.add_run(' shows after one day.')

   
    document.add_picture('10shows.jpeg', width=Cm(6.5), height=Cm(4.1))
    document.add_picture('PieChart.jpeg', width=Cm(6.5), height=Cm(4.1))
    document.add_picture('View_day.jpeg', width=Cm(6.5), height=Cm(4.1))

    document.save('Netflix_Analysis.docx')


if __name__ == '__main__':

    requirements()
    
    df, Count = data_cleaning()

    Netflix_time(df)
    left(df)

    views_by_day(df)
    ten(df)
    create_pie(Count)

    create_doc()
    os.startfile(os.path.abspath('Netflix_Analysis.docx'))
